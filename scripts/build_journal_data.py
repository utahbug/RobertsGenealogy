"""Build reader data and editable exports from authoritative journal Markdown."""

from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DATA_DIR = ROOT / "data" / "journals" / "rdr"
MANIFEST = DATA_DIR / "journal.json"
OUTPUT = DATA_DIR / "journal-data.json"
EXPORT_DIR = ROOT / "exports"
MARKDOWN_EXPORT = EXPORT_DIR / "RDR-Journal-Transcription.md"
DOCX_EXPORT = EXPORT_DIR / "RDR-Journal-Transcription.docx"
ALLOWED_STATUSES = {"untranscribed", "machine-draft", "needs-review", "reviewed"}
SECTION_KEYS = {
    "English transcription": "transcription",
    "Editorial annotations": "notes",
}

BODY_COLOR = RGBColor(40, 44, 47)
MUTED_COLOR = RGBColor(104, 110, 114)
TITLE_COLOR = RGBColor(41, 58, 67)
HEADING_BLUE = RGBColor(46, 116, 181)
HEADING_DARK_BLUE = RGBColor(31, 77, 120)


def parse_sections(path: Path) -> dict[str, str]:
    sections = {value: "" for value in SECTION_KEYS.values()}
    current = None
    collected: dict[str, list[str]] = {value: [] for value in SECTION_KEYS.values()}

    for line in path.read_text(encoding="utf-8").splitlines():
        if line.startswith("## "):
            current = SECTION_KEYS.get(line[3:].strip())
            continue
        if current:
            collected[current].append(line)

    for key, lines in collected.items():
        sections[key] = "\n".join(lines).strip()
    return sections


def set_style_font(style, name: str, size: float, color: RGBColor, bold: bool = False) -> None:
    style.font.name = name
    style.font.size = Pt(size)
    style.font.color.rgb = color
    style.font.bold = bold
    style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)


def set_run_font(run, name: str = "Calibri", size: float = 11, color: RGBColor = BODY_COLOR, bold: bool = False, italic: bool = False) -> None:
    run.font.name = name
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)


def add_page_number(paragraph) -> None:
    paragraph.add_run("Page ")
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = "PAGE"
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, text, end])
    set_run_font(run, size=9, color=MUTED_COLOR)


def page_heading(page: dict) -> str:
    manuscript = page.get("manuscriptPageNumber")
    if manuscript:
        return f"PDF page {page['pdfPageNumber']} - manuscript page {manuscript}"
    return f"PDF page {page['pdfPageNumber']}"


def add_text_block(document: Document, text: str) -> None:
    blocks = text.split("\n\n") if text else ["[Untranscribed]"]
    for block in blocks:
        paragraph = document.add_paragraph()
        lines = block.splitlines() or [""]
        for index, line in enumerate(lines):
            if index:
                paragraph.add_run().add_break()
            run = paragraph.add_run(line)
            set_run_font(run)


def build_reader_data(journal: dict) -> list[dict]:
    if journal.get("pdfPageCount") != len(journal.get("pages", [])):
        raise SystemExit("Journal manifest page count does not match its pages array")

    journal_pages = []
    for expected_number, page in enumerate(journal["pages"], start=1):
        if page.get("pdfPageNumber") != expected_number:
            raise SystemExit(f"Page sequence mismatch at PDF page {expected_number}")
        expected_id = f"journal-rdr-page-{expected_number:03d}"
        if page.get("id") != expected_id:
            raise SystemExit(f"Stable ID mismatch: expected {expected_id}, found {page.get('id')}")
        if page.get("transcriptionStatus") not in ALLOWED_STATUSES:
            raise SystemExit(f"Invalid transcription status for {expected_id}")

        transcription_path = ROOT / page["transcriptionFile"]
        image_path = ROOT / page["image"]
        if not transcription_path.exists():
            raise SystemExit(f"Missing transcription file: {transcription_path}")
        if not image_path.exists():
            raise SystemExit(f"Missing journal page image: {image_path}")

        sections = parse_sections(transcription_path)
        display_page = page.get("manuscriptPageNumber") or f"PDF {expected_number}"
        journal_pages.append(
            {
                "id": page["id"],
                "journalId": journal["id"],
                "journalTitle": journal["title"],
                "title": f"{journal['title']} - Page {display_page}",
                "pageNumber": expected_number,
                "pdfPageNumber": expected_number,
                "manuscriptPageNumber": page.get("manuscriptPageNumber"),
                "image": page["image"],
                "transcriptionFile": page["transcriptionFile"],
                "transcriptionStatus": page["transcriptionStatus"],
                "hand": page.get("hand"),
                "contentAuthor": page.get("contentAuthor"),
                "contentAuthorId": page.get("contentAuthorId"),
                "scribe": page.get("scribe"),
                "scribeId": page.get("scribeId"),
                "translator": page.get("translator"),
                "translatorId": page.get("translatorId"),
                "editor": page.get("editor"),
                "editorId": page.get("editorId"),
                "attributionStatus": page.get("attributionStatus"),
                "attributionBasis": page.get("attributionBasis"),
                "primaryWriter": journal["primaryWriter"],
                "transcriptionLabel": "English transcription",
                "searchLayer": page.get("searchLayer") or "English transcription",
                "notesLabel": "Editorial annotations",
                "transcription": sections["transcription"],
                "notes": sections["notes"],
                "contentContext": page.get("contentContext"),
                "people": page.get("people", []),
                "places": page.get("places", []),
                "dates": page.get("dates", []),
                "linkedEvents": page.get("linkedEvents", []),
                "editorialNotes": page.get("editorialNotes", []),
                "relatedJournalPassages": page.get("relatedJournalPassages", []),
            }
        )
    return journal_pages


def write_markdown_export(journal: dict, pages: list[dict]) -> None:
    lines = [
        f"# {journal['title']}",
        "",
        "Editable English transcription generated from the authoritative page Markdown files.",
        "",
        f"Primary writer: {journal['primaryWriter']}",
        f"Language: {journal['language']}",
        f"Source PDF: {journal['sourceFile']}",
        f"PDF pages: {journal['pdfPageCount']}",
        "",
        f"> {journal['openingContentNote']}",
        "",
    ]
    for page in pages:
        lines.extend(
            [
                f"## {page_heading(page)}",
                "",
                f"Stable ID: `{page['id']}`  ",
                f"Transcription status: `{page['transcriptionStatus']}`",
                "",
                "### English transcription",
                "",
                page["transcription"] or "[Untranscribed]",
                "",
                "### Editorial annotations",
                "",
                page["notes"] or "[No editorial annotations]",
                "",
            ]
        )
    MARKDOWN_EXPORT.write_text("\n".join(lines).rstrip() + "\n", encoding="utf-8")


def write_docx_export(journal: dict, pages: list[dict]) -> None:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = document.styles["Normal"]
    set_style_font(normal, "Calibri", 11, BODY_COLOR)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading1 = document.styles["Heading 1"]
    set_style_font(heading1, "Calibri", 16, HEADING_BLUE, True)
    heading1.paragraph_format.space_before = Pt(18)
    heading1.paragraph_format.space_after = Pt(10)
    heading1.paragraph_format.keep_with_next = True

    heading2 = document.styles["Heading 2"]
    set_style_font(heading2, "Calibri", 13, HEADING_BLUE, True)
    heading2.paragraph_format.space_before = Pt(14)
    heading2.paragraph_format.space_after = Pt(7)
    heading2.paragraph_format.keep_with_next = True

    heading3 = document.styles["Heading 3"]
    set_style_font(heading3, "Calibri", 12, HEADING_DARK_BLUE, True)
    heading3.paragraph_format.space_before = Pt(10)
    heading3.paragraph_format.space_after = Pt(5)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_run_font(header.add_run("Robert D. Roberts Journal | English Transcription"), size=9, color=MUTED_COLOR)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_page_number(footer)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(96)
    title.paragraph_format.space_after = Pt(8)
    set_run_font(title.add_run(journal["title"]), size=28, color=TITLE_COLOR, bold=True)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(28)
    set_run_font(subtitle.add_run("Editable English Transcription"), size=15, color=MUTED_COLOR)

    source = document.add_paragraph()
    source.alignment = WD_ALIGN_PARAGRAPH.CENTER
    source.paragraph_format.space_after = Pt(8)
    set_run_font(source.add_run(f"Source: {journal['sourceFile']} | {journal['pdfPageCount']} PDF pages"), size=10, color=MUTED_COLOR)

    status = document.add_paragraph()
    status.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(status.add_run("Working transcription: no page is marked reviewed; uncertainty notation is retained."), size=10, color=MUTED_COLOR, italic=True)

    document.add_page_break()

    context_heading = document.add_paragraph("Source and authorship context", style="Heading 1")
    context_heading.paragraph_format.space_before = Pt(0)
    add_text_block(document, journal["openingContentNote"])
    add_text_block(document, journal["authorshipNote"])

    for page in pages:
        document.add_paragraph(page_heading(page), style="Heading 1")
        metadata = document.add_paragraph()
        metadata.paragraph_format.space_after = Pt(7)
        set_run_font(
            metadata.add_run(f"Stable ID: {page['id']} | Transcription status: {page['transcriptionStatus']}"),
            size=9,
            color=MUTED_COLOR,
            italic=True,
        )
        if page.get("contentContext"):
            add_text_block(document, page["contentContext"])
        document.add_paragraph("English transcription", style="Heading 2")
        add_text_block(document, page["transcription"])
        document.add_paragraph("Editorial annotations", style="Heading 2")
        add_text_block(document, page["notes"])

    document.core_properties.title = "RDR Journal Transcription"
    document.core_properties.subject = "Editable English transcription of the Robert D. Roberts Journal"
    document.core_properties.author = "RobertsGenealogy"
    document.save(DOCX_EXPORT)


def main() -> None:
    journal = json.loads(MANIFEST.read_text(encoding="utf-8"))
    journal_pages = build_reader_data(journal)
    EXPORT_DIR.mkdir(parents=True, exist_ok=True)

    public_journal = {key: value for key, value in journal.items() if key != "pages"}
    OUTPUT.write_text(
        json.dumps({"journals": [public_journal], "journalPages": journal_pages}, indent=2, ensure_ascii=False) + "\n",
        encoding="utf-8",
    )
    write_markdown_export(journal, journal_pages)
    write_docx_export(journal, journal_pages)
    print(f"Built reader data and exports for {len(journal_pages)} journal pages")


if __name__ == "__main__":
    main()
